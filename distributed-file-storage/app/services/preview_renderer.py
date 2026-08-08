from pathlib import Path
from html import escape

from docx import Document
from openpyxl import load_workbook


# ============================================================
# EXCEL PREVIEW
# ============================================================

def render_xlsx(file_path: str) -> str:

    workbook = load_workbook(
        filename=file_path,
        read_only=True,
        data_only=True
    )

    html = """
    <!DOCTYPE html>
    <html>
    <head>
        <meta charset="UTF-8">
        <title>Spreadsheet Preview</title>

        <style>
            body {
                font-family: Arial, sans-serif;
                margin: 0;
                padding: 20px;
                background: #f5f7fa;
            }

            .container {
                background: white;
                padding: 20px;
                border-radius: 10px;
                overflow-x: auto;
            }

            h2 {
                margin-top: 0;
            }

            table {
                border-collapse: collapse;
                width: 100%;
                min-width: 600px;
            }

            th,
            td {
                border: 1px solid #ddd;
                padding: 8px 12px;
                text-align: left;
                white-space: nowrap;
            }

            th {
                background: #f1f3f5;
                font-weight: 600;
            }

            tr:nth-child(even) {
                background: #fafafa;
            }
        </style>
    </head>

    <body>
        <div class="container">
    """

    for worksheet in workbook.worksheets:

        html += f"""
        <h2>{escape(worksheet.title)}</h2>
        <table>
        """

        rows = worksheet.iter_rows(values_only=True)

        first_row = True

        for row in rows:

            # Skip completely empty rows
            if not any(value is not None for value in row):
                continue

            html += "<tr>"

            for value in row:

                value = "" if value is None else str(value)

                if first_row:
                    html += f"<th>{escape(value)}</th>"
                else:
                    html += f"<td>{escape(value)}</td>"

            html += "</tr>"

            first_row = False

        html += "</table>"

    html += """
        </div>
    </body>
    </html>
    """

    workbook.close()

    return html


# ============================================================
# DOCX PREVIEW
# ============================================================

def render_docx(file_path: str) -> str:

    document = Document(file_path)

    html = """
    <!DOCTYPE html>
    <html>
    <head>
        <meta charset="UTF-8">

        <style>

            body {
                background: #f5f7fa;
                font-family: Arial, sans-serif;
                margin: 0;
                padding: 40px;
            }

            .document {
                background: white;
                max-width: 900px;
                margin: auto;
                padding: 60px;
                box-shadow: 0 2px 10px rgba(0,0,0,0.08);
            }

            p {
                line-height: 1.7;
                margin-bottom: 12px;
            }

            table {
                border-collapse: collapse;
                width: 100%;
                margin: 20px 0;
            }

            td,
            th {
                border: 1px solid #ddd;
                padding: 8px;
            }

        </style>

    </head>

    <body>

        <div class="document">
    """

    # Paragraphs
    for paragraph in document.paragraphs:

        text = paragraph.text.strip()

        if not text:
            continue

        html += f"<p>{escape(text)}</p>"

    # Tables
    for table in document.tables:

        html += "<table>"

        for row in table.rows:

            html += "<tr>"

            for cell in row.cells:

                text = cell.text.strip()

                html += f"<td>{escape(text)}</td>"

            html += "</tr>"

        html += "</table>"

    html += """
        </div>

    </body>
    </html>
    """

    return html